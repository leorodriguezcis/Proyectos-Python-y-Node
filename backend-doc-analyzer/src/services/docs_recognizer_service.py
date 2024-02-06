from fastapi import UploadFile
from src.dto.requests.document import Document
import docx
from src.decorators.map_exceptions import exceptions_mapped, map_exceptions
from src.exceptions.task.service_unavailable_exception import (
    ServiceUnavailableException,
)


@map_exceptions(
    target_exception=ServiceUnavailableException(detail="Docs recognizer Error")
)
class DocsRecognizerService:
    @exceptions_mapped
    def analyze_document(self, document: UploadFile):
        try:
            doc = docx.Document(document.file._file)
        except Exception:
            doc = docx.Document(document.file)

        fullText = ""

        for paragraph in doc.paragraphs:
            fullText += paragraph.text + "\n"

        return fullText

    @staticmethod
    def get_instance() -> "DocsRecognizerService":
        return DocsRecognizerService()
